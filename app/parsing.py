# app/parsing.py - Advanced resume parsing with comprehensive entity extraction and EDA
from __future__ import annotations
import hashlib
import re
from collections import Counter
from collections import defaultdict
from pathlib import Path
from typing import Any
from typing import Dict
from typing import List
from typing import Tuple

import dateparser
import phonenumbers
import pdfplumber
from docx import Document as DocxDocument
from pypdf import PdfReader

# Advanced resume parsing libraries
try:
    from pyresparser import ResumeParser as PyResParser
    _HAS_PYRESPARSER = True
except Exception:
    _HAS_PYRESPARSER = False

try:
    import layoutparser as lp
    _HAS_LAYOUTPARSER = True
except Exception:
    _HAS_LAYOUTPARSER = False

# Advanced NLP libraries
try:
    import spacy
    # Try to load transformer model first, fallback to standard
    try:
        _NLP = spacy.load("en_core_web_trf")  # Transformer-based model
    except OSError:
        _NLP = spacy.load("en_core_web_sm")  # Standard model
except Exception:
    _NLP = None

try:
    from transformers import pipeline
    # Load advanced NER model for better entity extraction
    _NER_PIPELINE = pipeline("ner", 
                            model="dbmdz/bert-large-cased-finetuned-conll03-english",
                            aggregation_strategy="simple")
except Exception:
    _NER_PIPELINE = None

try:
    import textstat
    _HAS_TEXTSTAT = True
except ImportError:
    _HAS_TEXTSTAT = False

# ---------- Config ----------
ALLOWED = {".pdf", ".doc", ".docx"}
MAX_FILES_PER_SESSION = 50

# Social link patterns (extendable)
SOCIAL_PATTERNS = [
    r"https?://(www\.)?linkedin\.com/[A-Za-z0-9_/\-?=%.]+",
    r"https?://(www\.)?github\.com/[A-Za-z0-9_\-./]+",
    r"https?://(www\.)?(x|twitter)\.com/[A-Za-z0-9_\-./]+",
    r"https?://(www\.)?kaggle\.com/[A-Za-z0-9_\-./]+",
    r"https?://(www\.)?medium\.com/[A-Za-z0-9_\-./]+",
    r"https?://(www\.)?personal\.[A-Za-z0-9_\-./]+",
    r"https?://[A-Za-z0-9_\-./]+\.me(/[A-Za-z0-9_\-./]+)?",
    r"mailto:[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
]
SOCIAL = re.compile("|".join(SOCIAL_PATTERNS), re.IGNORECASE)

EMAIL = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
URL = re.compile(r"https?://[^\s)>\]]+", re.IGNORECASE)
PHONE = re.compile(r"(\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})")
DATE_PATTERNS = [
    r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b",
    r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    r"\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b",
    r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b"
]

# Comprehensive skills database with categories
SKILLS_DATABASE = {
    "programming": {
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", 
        "swift", "kotlin", "scala", "r", "matlab", "perl", "shell", "bash", "powershell"
    },
    "web_development": {
        "html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask", 
        "fastapi", "spring", "laravel", "rails", "asp.net", "jquery", "bootstrap", "sass", "less"
    },
    "databases": {
        "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra", "oracle", 
        "sql server", "sqlite", "dynamodb", "neo4j", "influxdb", "couchdb"
    },
    "cloud_platforms": {
        "aws", "azure", "gcp", "google cloud", "heroku", "digitalocean", "linode", "vultr"
    },
    "devops": {
        "docker", "kubernetes", "jenkins", "gitlab ci", "github actions", "terraform", "ansible", 
        "puppet", "chef", "vagrant", "prometheus", "grafana", "elk stack", "splunk"
    },
    "data_science": {
        "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "xgboost", 
        "lightgbm", "spark", "hadoop", "tableau", "power bi", "jupyter", "r studio"
    },
    "mobile": {
        "android", "ios", "react native", "flutter", "ionic", "xamarin", "cordova"
    },
    "testing": {
        "selenium", "cypress", "jest", "mocha", "pytest", "junit", "testng", "cucumber"
    },
    "soft_skills": {
        "leadership", "communication", "teamwork", "problem solving", "project management", 
        "agile", "scrum", "kanban", "mentoring", "training", "presentation", "negotiation"
    }
}

# ---------- Helpers ----------
def file_sha256(p: Path) -> str:
    h = hashlib.sha256()
    with open(p, "rb") as f:
        for b in iter(lambda: f.read(1024*1024), b""):
            h.update(b)
    return h.hexdigest()

def _extract_names(text: str) -> List[str]:
    """Extract names using multiple strategies for comprehensive coverage"""
    names = []
    
    # Strategy 1: Look for common name patterns in resume headers
    name_patterns = [
        r"^([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*$",  # First line names
        r"(?:Name|NAME):\s*([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)",  # "Name: John Doe"
        r"(?:Resume of|CV of)\s+([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)",  # "Resume of John Doe"
    ]
    
    # Focus on the first few lines for names
    first_lines = text.split('\n')[:5]
    for line in first_lines:
        line = line.strip()
        for pattern in name_patterns:
            matches = re.findall(pattern, line, re.MULTILINE | re.IGNORECASE)
            names.extend(matches)
    
    # Strategy 2: Look for standalone names in the first few lines
    for line in first_lines:
        line = line.strip()
        # Look for lines that might be names (2-3 words, properly capitalized)
        if re.match(r'^[A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$', line):
            # Filter out common false positives
            if len(line.split()) <= 3 and not any(word.lower() in line.lower() for word in 
                                                 ["resume", "cv", "curriculum", "vitae", "portfolio", "engineer", "developer", "manager", "analyst", "summary", "professional"]):
                names.append(line)
    
    # Strategy 3: Use spaCy NER if available
    if _NLP:
        try:
            # Only process first 1000 chars for performance
            doc = _NLP(text[:1000])  
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                    # Filter out common false positives
                    if not any(word.lower() in ent.text.lower() for word in 
                              ["company", "university", "college", "school", "inc", "corp", "ltd", "engineer", "developer"]):
                        # Only add if it looks like a proper name (capitalized words)
                        if all(word[0].isupper() for word in ent.text.split() if word.isalpha()):
                            names.append(ent.text)
        except Exception:
            pass
    
    # Strategy 4: Extract names from email addresses (e.g., john.doe@email.com)
    email_names = []
    for email in EMAIL.findall(text):
        local_part = email.split('@')[0]
        # Convert john.doe or john_doe to John Doe
        if '.' in local_part or '_' in local_part:
            name_parts = re.split(r'[._]', local_part)
            if len(name_parts) >= 2 and all(part.isalpha() and len(part) > 1 for part in name_parts):
                formatted_name = ' '.join(part.capitalize() for part in name_parts)
                email_names.append(formatted_name)
    names.extend(email_names)
    
    return names

def _extract_skills_with_context(text: str) -> List[str]:
    """Extract skills with enhanced context awareness and semantic grouping"""
    text_lower = text.lower()
    found_skills = []
    
    # Use existing skill detection as base
    base_skills = _skills_from_text(text)
    found_skills.extend(base_skills)
    
    # Enhanced skill detection with categories
    skill_categories = {
        "programming_languages": [
            "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", 
            "swift", "kotlin", "scala", "r", "matlab", "perl", "shell scripting", "bash", "powershell",
            "objective-c", "dart", "julia", "haskell", "clojure", "erlang", "elixir"
        ],
        "web_technologies": [
            "html", "css", "react", "angular", "vue", "vue.js", "node.js", "express", "django", 
            "flask", "fastapi", "spring boot", "laravel", "rails", "asp.net", "jquery", 
            "bootstrap", "sass", "less", "webpack", "babel", "next.js", "nuxt.js", "svelte"
        ],
        "databases": [
            "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra", "oracle", 
            "sql server", "sqlite", "dynamodb", "neo4j", "influxdb", "couchdb", "mariadb",
            "firestore", "cosmos db", "bigquery", "snowflake", "databricks"
        ],
        "cloud_platforms": [
            "aws", "azure", "gcp", "google cloud", "heroku", "digitalocean", "linode", "vultr",
            "cloudflare", "netlify", "vercel", "firebase", "supabase", "planetscale"
        ],
        "devops_tools": [
            "docker", "kubernetes", "jenkins", "gitlab ci", "github actions", "terraform", 
            "ansible", "puppet", "chef", "vagrant", "prometheus", "grafana", "elk stack", 
            "splunk", "datadog", "new relic", "circleci", "travis ci", "helm", "istio"
        ],
        "data_science": [
            "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "xgboost", 
            "lightgbm", "spark", "hadoop", "tableau", "power bi", "jupyter", "r studio",
            "matplotlib", "seaborn", "plotly", "dask", "airflow", "mlflow", "kubeflow"
        ],
        "testing_frameworks": [
            "selenium", "cypress", "jest", "mocha", "pytest", "junit", "testng", "cucumber",
            "postman", "insomnia", "k6", "locust", "jmeter", "robot framework"
        ],
        "soft_skills": [
            "leadership", "communication", "teamwork", "problem solving", "project management", 
            "agile", "scrum", "kanban", "mentoring", "training", "presentation", "negotiation",
            "time management", "critical thinking", "adaptability", "collaboration"
        ]
    }
    
    # Enhanced pattern matching with context
    for category, skills in skill_categories.items():
        for skill in skills:
            # Create flexible pattern for skill matching
            skill_pattern = r'\b' + re.escape(skill).replace(r'\ ', r'[\s\-\_]*') + r'\b'
            if re.search(skill_pattern, text_lower):
                found_skills.append(skill)
    
    # Look for skills in specific sections
    skill_sections = re.findall(r'(?:SKILLS?|TECHNICAL SKILLS?|TECHNOLOGIES?)[:\s]*\n(.*?)(?:\n\s*\n|\n[A-Z]|\Z)', 
                               text, re.IGNORECASE | re.DOTALL)
    
    for section in skill_sections:
        # Extract comma-separated or bullet-pointed skills
        skill_items = re.findall(r'[•\-\*]?\s*([A-Za-z0-9\+\#\.\s]{2,30})', section)
        for item in skill_items:
            item = item.strip().lower()
            if item and len(item.split()) <= 3:  # Avoid long phrases
                # Check if it's a known technology
                for category, skills in skill_categories.items():
                    if item in [s.lower() for s in skills]:
                        found_skills.append(item)
    
    return found_skills

def _extract_enhanced_social_links(text: str) -> List[str]:
    """Enhanced social link detection with better pattern matching"""
    social_links = []
    
    # Enhanced social media patterns
    enhanced_patterns = [
        # LinkedIn variations
        r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-\.]+/?',
        r'(?:https?://)?(?:www\.)?linkedin\.com/pub/[A-Za-z0-9\-\.\/]+/?',
        r'(?:https?://)?(?:www\.)?linkedin\.com/profile/view\?id=[A-Za-z0-9\-\.]+',
        
        # GitHub variations
        r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9\-\.]+/?',
        r'(?:https?://)?(?:www\.)?github\.io/[A-Za-z0-9\-\.\/]+/?',
        
        # Twitter/X variations
        r'(?:https?://)?(?:www\.)?(?:twitter|x)\.com/[A-Za-z0-9_]+/?',
        
        # Other professional platforms
        r'(?:https?://)?(?:www\.)?stackoverflow\.com/users/[A-Za-z0-9\-\/]+',
        r'(?:https?://)?(?:www\.)?kaggle\.com/[A-Za-z0-9\-]+/?',
        r'(?:https?://)?(?:www\.)?medium\.com/@?[A-Za-z0-9\-\.]+/?',
        r'(?:https?://)?(?:www\.)?dev\.to/[A-Za-z0-9\-]+/?',
        r'(?:https?://)?(?:www\.)?codepen\.io/[A-Za-z0-9\-]+/?',
        r'(?:https?://)?(?:www\.)?behance\.net/[A-Za-z0-9\-]+/?',
        r'(?:https?://)?(?:www\.)?dribbble\.com/[A-Za-z0-9\-]+/?',
        
        # Personal websites and portfolios
        r'(?:https?://)?[A-Za-z0-9\-]+\.(?:dev|me|io|com|net|org)/?',
        r'(?:https?://)?(?:www\.)?[A-Za-z0-9\-]+\.(?:portfolio|website)\..*',
    ]
    
    for pattern in enhanced_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        social_links.extend(matches)
    
    # Look for email-style social mentions (e.g., "LinkedIn: john.doe")
    social_mention_patterns = [
        r'LinkedIn:\s*([A-Za-z0-9\.\-_]+)',
        r'GitHub:\s*([A-Za-z0-9\.\-_]+)', 
        r'Twitter:\s*@?([A-Za-z0-9_]+)',
        r'Portfolio:\s*((?:https?://)?[A-Za-z0-9\-\.]+\.[A-Za-z]{2,})',
    ]
    
    for pattern in social_mention_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if 'linkedin' in pattern.lower():
                social_links.append(f"https://linkedin.com/in/{match}")
            elif 'github' in pattern.lower():
                social_links.append(f"https://github.com/{match}")
            elif 'twitter' in pattern.lower():
                social_links.append(f"https://twitter.com/{match}")
            else:
                social_links.append(match)
    
    return social_links

def _extract_organizations(text: str) -> List[str]:
    """Extract organization names using NLP and patterns"""
    organizations = []
    
    # Use spaCy NER if available
    if _NLP:
        try:
            doc = _NLP(text)
            for ent in doc.ents:
                if ent.label_ in ["ORG"]:
                    # Filter out common false positives
                    if len(ent.text) > 2 and not ent.text.lower() in ["resume", "cv", "email", "phone"]:
                        organizations.append(ent.text)
        except Exception:
            pass
    
    # Pattern-based organization extraction
    org_patterns = [
        r'(?:at|@)\s+([A-Z][A-Za-z\s&\.,]+(?:Inc|Corp|LLC|Ltd|Company|Co|University|College|School))',
        r'([A-Z][A-Za-z\s&\.,]+(?:Inc|Corp|LLC|Ltd|Company|Co|Technologies|Tech|Solutions|Systems))',
        r'([A-Z][A-Za-z\s&\.,]+(?:University|College|School|Institute))',
    ]
    
    for pattern in org_patterns:
        matches = re.findall(pattern, text)
        organizations.extend(matches)
    
    return organizations

def _extract_locations(text: str) -> List[str]:
    """Extract location names using NLP and patterns"""
    locations = []
    
    # Use spaCy NER if available
    if _NLP:
        try:
            doc = _NLP(text)
            for ent in doc.ents:
                if ent.label_ in ["GPE", "LOC"]:  # Geopolitical entities and locations
                    locations.append(ent.text)
        except Exception:
            pass
    
    # Pattern-based location extraction (common formats in resumes)
    location_patterns = [
        r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, State (US format)
        r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, Country
        r'([A-Z][a-z]+\s+[A-Z][a-z]+,\s*[A-Z]{2})',  # City Name, State
    ]
    
    for pattern in location_patterns:
        matches = re.findall(pattern, text)
        locations.extend(matches)
    
    return locations

def _extract_experience_years(text: str) -> List[str]:
    """Extract years of experience from text"""
    experience_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)',
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:in|with|of)',
        r'(\d+)\+?\s*(?:year|yr)\s*(?:experience|exp)',
        r'(?:over|more than|about|approximately)\s*(\d+)\s*(?:years?|yrs?)',
        r'(\d+)\s*to\s*(\d+)\s*(?:years?|yrs?)',
    ]
    
    experience_years = []
    for pattern in experience_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                experience_years.extend(match)
            else:
                experience_years.append(match)
    
    return experience_years

def _extract_certifications(text: str) -> List[str]:
    """Extract certifications and professional credentials"""
    cert_patterns = [
        r'(?:AWS|Amazon)\s*(?:Certified|Certification)?\s*([A-Za-z\s\-]+)',
        r'(?:Microsoft|Azure)\s*(?:Certified|Certification)?\s*([A-Za-z\s\-]+)',
        r'(?:Google|GCP)\s*(?:Certified|Certification)?\s*([A-Za-z\s\-]+)',
        r'(?:Cisco|CCNA|CCNP|CCIE)',
        r'(?:PMP|PMI|Project Management Professional)',
        r'(?:CISSP|CISM|CISA|CEH)',
        r'(?:Scrum Master|Agile|SAFe)',
        r'([A-Z]{2,})\s*(?:Certified|Certification)',
    ]
    
    certifications = []
    for pattern in cert_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            if isinstance(matches[0], str):
                certifications.extend(matches)
            else:
                certifications.extend([m for m in matches if m])
    
    return certifications

def _extract_education(text: str) -> List[str]:
    """Extract education information"""
    education = []
    
    # Degree patterns
    degree_patterns = [
        r'(Bachelor(?:\'s)?|BS|BA|B\.S\.|B\.A\.)\s*(?:of|in|degree)?\s*([A-Za-z\s]+)',
        r'(Master(?:\'s)?|MS|MA|M\.S\.|M\.A\.)\s*(?:of|in|degree)?\s*([A-Za-z\s]+)',
        r'(PhD|Ph\.D\.|Doctorate|Doctoral)\s*(?:of|in|degree)?\s*([A-Za-z\s]+)',
        r'(Associate(?:\'s)?|AS|AA|A\.S\.|A\.A\.)\s*(?:of|in|degree)?\s*([A-Za-z\s]+)',
    ]
    
    for pattern in degree_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                education.append(' '.join(match).strip())
            else:
                education.append(match)
    
    return education

def _calculate_readability_stats(text: str) -> Dict[str, Any]:
    """Calculate readability statistics for the text"""
    if not _HAS_TEXTSTAT:
        return {}
    
    try:
        stats = {
            "flesch_reading_ease": textstat.flesch_reading_ease(text),
            "flesch_kincaid_grade": textstat.flesch_kincaid_grade(text),
            "gunning_fog": textstat.gunning_fog(text),
            "automated_readability_index": textstat.automated_readability_index(text),
            "coleman_liau_index": textstat.coleman_liau_index(text),
            "reading_time_minutes": textstat.reading_time(text),
        }
        return stats
    except Exception:
        return {}

def _extract_comprehensive_entities(text: str) -> Dict[str, Any]:

    # --- Ensemble entity extraction ---
    entities = {
        "names": [],
        "emails": [],
        "phones": [],
        "dates": [],
        "organizations": [],
        "locations": [],
        "skills": [],
        "certifications": [],
        "education": [],
        "experience_years": [],
        "readability_stats": {}
    }

    # 1. Basic regex patterns
    entities["emails"] = EMAIL.findall(text)
    entities["phones"] = PHONE.findall(text)
    
    # 2. Extract names using multiple strategies
    entities["names"] = _extract_names(text)
    
    # 3. Extract skills with context awareness
    entities["skills"] = _extract_skills_with_context(text)
    
    # 4. Extract social links with enhanced detection
    social_urls = _extract_enhanced_social_links(text)
    
    # 5. Extract dates using enhanced patterns
    for pattern in DATE_PATTERNS:
        dates = re.findall(pattern, text, re.IGNORECASE)
        entities["dates"].extend(dates)
    
    # 6. Extract organizations using NLP if available
    entities["organizations"] = _extract_organizations(text)
    
    # 7. Extract locations using NLP if available
    entities["locations"] = _extract_locations(text)
    
    # 8. Extract experience years
    entities["experience_years"] = _extract_experience_years(text)
    
    # 9. Extract certifications
    entities["certifications"] = _extract_certifications(text)
    
    # 10. Extract education information
    entities["education"] = _extract_education(text)
    
    # 11. Calculate readability statistics
    if _HAS_TEXTSTAT:
        entities["readability_stats"] = _calculate_readability_stats(text)

    # 12. Deduplicate and clean
    for key in entities:
        if isinstance(entities[key], list):
            # Remove empty strings and duplicates while preserving order
            seen = set()
            cleaned = []
            for item in entities[key]:
                if item and item not in seen:
                    seen.add(item)
                    cleaned.append(item)
            entities[key] = cleaned

    return entities

def _extract_docx_hyperlinks(doc: DocxDocument) -> List[str]:
    # Collect explicit hyperlink targets from DOCX rels
    urls = []
    rels = getattr(doc.part, "rels", {})
    for r in rels.values():
        try:
            if "hyperlink" in r.reltype and r.target_ref:
                urls.append(str(r.target_ref))
        except Exception:
            continue
    return urls

def _parse_pdf(p: Path) -> Tuple[str, List[str], Dict[int, List[str]]]:
    """Return full_text, url_list, page2urls"""
    text_parts = []
    page2urls: Dict[int, List[str]] = {}
    urls = []

    # Primary: pdfplumber (layout aware)
    try:
        with pdfplumber.open(str(p)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
                    # Extract URLs from text
                    page_urls = URL.findall(t)
                    # Try PDF annotations (link rectangles)
                    try:
                        raw_annots = page.annotations or []
                        for a in raw_annots:
                            uri = a.get("uri") or a.get("A", {}).get("URI")
                            if uri:
                                page_urls.append(uri)
                    except Exception:
                        pass
                    if page_urls:
                        page2urls[i+1] = list(sorted(set(page_urls)))
                        urls.extend(page_urls)
    except Exception:
        # Fallback: PyPDF
        rd = PdfReader(str(p))
        for i, pg in enumerate(rd.pages):
            t = pg.extract_text() or ""
            if t.strip():
                text_parts.append(t)
                page_urls = URL.findall(t)
                if page_urls:
                    page2urls[i+1] = list(sorted(set(page_urls)))
                    urls.extend(page_urls)

    return "\n".join(text_parts), sorted(set(urls)), page2urls

def _parse_docx(p: Path) -> Tuple[str, List[str]]:
    """Enhanced DOCX parsing with better text extraction"""
    d = DocxDocument(str(p))
    paragraphs = [para.text for para in d.paragraphs if para.text.strip()]
    
    # Extract table content more intelligently
    tables = []
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                tables.append(" | ".join(cells))
    
    # Extract headers and footers
    headers_footers = []
    try:
        for section in d.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        headers_footers.append(para.text.strip())
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        headers_footers.append(para.text.strip())
    except Exception:
        pass
    
    full = "\n".join(paragraphs + tables + headers_footers)
    urls = sorted(set(URL.findall(full) + _extract_docx_hyperlinks(d)))
    return full, urls

# ---------- Adaptive EDA-informed chunking ----------
def _adaptive_params(total_bytes: int, file_count: int) -> Dict[str, int]:
    """
    Dynamically scale chunk size & overlap based on corpus characteristics.
    
    Factors considered:
    - Total content size (larger content needs bigger chunks for context)
    - Number of files (more files might need smaller chunks for precision)
    - Content density estimation
    """
    mb = max(1, total_bytes // (1024*1024))
    
    # Base parameters optimized for resume processing
    base_chunk = 512  # Optimal for resume sections
    base_overlap = 64  # Good for maintaining context
    
    # Scale based on content size with smart thresholds
    if mb <= 1:  # Small files (typical single resume)
        chunk_multiplier = 1.0
        overlap_multiplier = 1.0
    elif mb <= 5:  # Medium files (multiple resumes or detailed resumes)
        chunk_multiplier = 1.2
        overlap_multiplier = 1.1
    elif mb <= 20:  # Large files (portfolios, detailed CVs)
        chunk_multiplier = 1.5
        overlap_multiplier = 1.2
    else:  # Very large files (document collections)
        chunk_multiplier = 2.0
        overlap_multiplier = 1.3
    
    # Adjust for file count - more files might need smaller chunks for precision
    file_factor = min(1.5, 1.0 + (file_count - 1) * 0.1)
    
    # Calculate final parameters
    chunk = int(base_chunk * chunk_multiplier * file_factor)
    overlap = int(base_overlap * overlap_multiplier)
    
    # Apply reasonable bounds
    chunk = min(2048, max(256, chunk))  # Between 256 and 2048 chars
    overlap = min(chunk // 3, max(32, overlap))  # Overlap shouldn't exceed 1/3 of chunk
    
    return {"chunk": chunk, "overlap": overlap}

def _split_overlap(text: str, chunk: int, overlap: int) -> List[str]:
    """Enhanced chunking with semantic boundary awareness"""
    text = text.replace("\r", "\n")
    
    # Try to split at natural boundaries when possible
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        # If paragraph is too long, split it using original method
        if len(paragraph) > chunk:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            
            # Split long paragraph with overlap
            para_chunks = _split_text_with_overlap(paragraph, chunk, overlap)
            chunks.extend(para_chunks)
        else:
            # Try to add paragraph to current chunk
            if len(current_chunk) + len(paragraph) + 2 <= chunk:
                current_chunk += ("\n\n" if current_chunk else "") + paragraph
            else:
                # Current chunk is full, start new one
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph
    
    # Add remaining chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # If no semantic splitting worked, fall back to character-based splitting
    if not chunks or (len(chunks) == 1 and len(chunks[0]) > chunk * 1.5):
        return _split_text_with_overlap(text, chunk, overlap)
    
    return chunks

def _split_text_with_overlap(text: str, chunk: int, overlap: int) -> List[str]:
    """Original character-based splitting with overlap"""
    out = []
    i = 0
    n = len(text)
    
    while i < n:
        j = min(i + chunk, n)
        seg = text[i:j].strip()
        if seg:
            out.append(seg)
        if j >= n:
            break
        i = max(0, j - overlap)
    
    return out

# ---------- Skills and entities ----------
EXTENDED_SKILLS = {
    # Programming Languages
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "rust", "ruby", "php",
    "swift", "kotlin", "scala", "r", "matlab", "perl", "shell", "bash", "powershell", "sql",
    "objective-c", "dart", "julia", "haskell", "clojure", "erlang", "elixir", "lua", "vba",
    
    # Web Technologies
    "html", "css", "react", "angular", "vue", "vue.js", "node", "node.js", "express", "flask", 
    "django", "fastapi", "spring", "spring boot", "laravel", "rails", "asp.net", "jquery",
    "bootstrap", "sass", "less", "webpack", "babel", "next.js", "nuxt.js", "svelte", "ember",
    
    # Databases
    "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra", "oracle", 
    "sql server", "sqlite", "dynamodb", "neo4j", "influxdb", "couchdb", "mariadb",
    "firestore", "cosmos db", "bigquery", "snowflake", "databricks",
    
    # Cloud Platforms & DevOps
    "aws", "azure", "gcp", "google cloud", "heroku", "digitalocean", "docker", "kubernetes",
    "jenkins", "gitlab ci", "github actions", "terraform", "ansible", "puppet", "chef",
    "vagrant", "prometheus", "grafana", "elk stack", "splunk", "datadog", "helm", "istio",
    
    # Data Science & ML
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "xgboost", 
    "lightgbm", "spark", "hadoop", "tableau", "power bi", "jupyter", "matplotlib",
    "seaborn", "plotly", "dask", "airflow", "mlflow", "kubeflow", "nlp", "ml", "ai",
    
    # Testing & QA
    "selenium", "cypress", "jest", "mocha", "pytest", "junit", "testng", "cucumber",
    "postman", "insomnia", "k6", "locust", "jmeter", "robot framework",
    
    # Mobile Development
    "android", "ios", "react native", "flutter", "ionic", "xamarin", "cordova",
    
    # Version Control & Tools
    "git", "svn", "mercurial", "bitbucket", "github", "gitlab", "jira", "confluence",
    "slack", "teams", "notion", "trello", "asana", "monday",
    
    # Operating Systems
    "linux", "windows", "macos", "ubuntu", "centos", "redhat", "debian", "unix",
    
    # Soft Skills & Methodologies
    "agile", "scrum", "kanban", "leadership", "communication", "teamwork", "mentoring",
    "project management", "problem solving", "critical thinking", "time management",
    "presentation", "negotiation", "training", "collaboration", "adaptability",
    
    # Business & Analytics
    "excel", "powerpoint", "word", "outlook", "sharepoint", "salesforce", "hubspot",
    "google analytics", "adobe analytics", "mixpanel", "amplitude",
    
    # Design & Creative
    "photoshop", "illustrator", "indesign", "sketch", "figma", "adobe xd", "canva",
    "after effects", "premiere pro", "blender", "maya", "3ds max"
}

def _skills_from_text(text: str) -> List[str]:
    """Enhanced skill extraction with better pattern matching"""
    t = text.lower()
    found = []
    
    for s in EXTENDED_SKILLS:
        # Create flexible pattern that handles variations
        skill_words = s.split()
        if len(skill_words) == 1:
            # Single word skill - use word boundary
            pat = r"\b" + re.escape(s) + r"\b"
        else:
            # Multi-word skill - allow flexible spacing and separators
            escaped_words = [re.escape(word) for word in skill_words]
            pat = r"\b" + r"[\s\-\_\.]*".join(escaped_words) + r"\b"
        
        if re.search(pat, t, re.IGNORECASE):
            found.append(s)
            
    return sorted(set(found))

# ---------- Public API ----------
def _perform_advanced_eda(text: str, entities: Dict[str, Any]) -> Dict[str, Any]:
    """Perform advanced exploratory data analysis on resume text with semantic patterns"""
    eda_results = {
        "text_statistics": {},
        "keyword_density": {},
        "semantic_themes": [],
        "experience_indicators": {},
        "skill_categories": {},
        "context_patterns": {},
        "professional_maturity": {},
        "semantic_sections": {},
        "sentiment_analysis": {}
    }
    
    # Basic text statistics
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    
    eda_results["text_statistics"] = {
        "total_words": len(words),
        "total_sentences": len([s for s in sentences if s.strip()]),
        "avg_words_per_sentence": len(words) / max(1, len(sentences)),
        "unique_words": len(set(word.lower() for word in words if word.isalpha())),
        "lexical_diversity": len(set(word.lower() for word in words if word.isalpha())) / max(1, len(words))
    }
    
    # Enhanced keyword density analysis
    word_freq = Counter(word.lower() for word in words if word.isalpha() and len(word) > 3)
    total_words = len(words)
    for word, count in word_freq.most_common(20):
        eda_results["keyword_density"][word] = count / total_words
    
    # Semantic section detection
    eda_results["semantic_sections"] = _detect_semantic_sections(text)
    
    # Enhanced experience indicators with context
    text_lower = text.lower()
    
    # Leadership and management indicators
    leadership_patterns = [
        r'(?:led|leading|lead)\s+(?:a\s+)?(?:team|group|project|initiative)',
        r'(?:managed|managing|manage)\s+(?:a\s+)?(?:team|group|department|project)',
        r'(?:supervised|supervising|supervise)\s+\d+',
        r'(?:coordinated|coordinating|coordinate)\s+(?:with|between|across)',
        r'(?:mentored|mentoring|mentor)\s+(?:junior|new|team)',
    ]
    
    technical_patterns = [
        r'(?:developed|developing|develop)\s+(?:a\s+)?(?:system|application|solution|feature)',
        r'(?:implemented|implementing|implement)\s+(?:a\s+)?(?:system|solution|algorithm)',
        r'(?:designed|designing|design)\s+(?:and\s+)?(?:implemented|developed|built)',
        r'(?:optimized|optimizing|optimize)\s+(?:performance|system|algorithm)',
        r'(?:built|building|build)\s+(?:from\s+scratch|a\s+)?(?:system|application)',
    ]
    
    collaboration_patterns = [
        r'(?:collaborated|collaborating|collaborate)\s+(?:with|across|between)',
        r'(?:worked\s+closely|working\s+closely)\s+with',
        r'(?:partnered|partnering|partner)\s+with',
        r'(?:cross-functional|interdisciplinary)\s+(?:team|collaboration)',
    ]
    
    achievement_patterns = [
        r'(?:increased|improved|reduced|achieved|delivered)\s+.*?(?:\d+%|\d+x)',
        r'(?:successfully|successfully\s+delivered|delivered)\s+.*?(?:on\s+time|ahead\s+of\s+schedule)',
        r'(?:exceeded|surpassed)\s+(?:expectations|targets|goals)',
        r'(?:awarded|recognized|selected)\s+(?:for|as)',
    ]
    
    eda_results["experience_indicators"] = {
        "leadership_score": len([p for p in leadership_patterns if re.search(p, text_lower, re.IGNORECASE)]),
        "technical_score": len([p for p in technical_patterns if re.search(p, text_lower, re.IGNORECASE)]),
        "collaboration_score": len([p for p in collaboration_patterns if re.search(p, text_lower, re.IGNORECASE)]),
        "achievement_score": len([p for p in achievement_patterns if re.search(p, text_lower, re.IGNORECASE)]),
    }
    
    # Context-aware skill categorization
    eda_results["skill_categories"] = _categorize_skills_with_context(text, entities.get("skills", []))
    
    # Professional maturity indicators
    eda_results["professional_maturity"] = _assess_professional_maturity(text, entities)
    
    # Context patterns for better semantic understanding
    eda_results["context_patterns"] = _extract_context_patterns(text)
    
    return eda_results

def _detect_semantic_sections(text: str) -> Dict[str, List[str]]:
    """Detect and categorize semantic sections in resume"""
    sections = {
        "contact": [],
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
        "achievements": []
    }
    
    # Section headers patterns
    section_patterns = {
        "contact": [r'contact(?:\s+(?:information|info|details))?', r'personal(?:\s+details)?'],
        "summary": [r'(?:professional\s+)?summary', r'profile', r'overview', r'objective', r'career\s+objective'],
        "experience": [r'(?:work\s+|professional\s+)?experience', r'employment(?:\s+history)?', r'career(?:\s+history)?'],
        "education": [r'education(?:al\s+background)?', r'academic(?:\s+background)?', r'qualifications'],
        "skills": [r'(?:technical\s+)?skills?', r'competencies', r'technologies', r'expertise'],
        "projects": [r'projects?', r'portfolio', r'selected\s+work'],
        "certifications": [r'certifications?', r'licenses?', r'credentials'],
        "achievements": [r'achievements?', r'accomplishments?', r'awards?', r'honors?']
    }
    
    lines = text.split('\n')
    current_section = None
    
    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            continue
            
        # Check if line is a section header
        for section_name, patterns in section_patterns.items():
            if any(re.search(rf'^{pattern}:?\s*$', line_clean) for pattern in patterns):
                current_section = section_name
                break
        
        # Add content to current section
        if current_section and line.strip():
            sections[current_section].append(line.strip())
    
    return sections

def _categorize_skills_with_context(text: str, skills: List[str]) -> Dict[str, List[str]]:
    """Categorize skills based on context where they appear"""
    skill_contexts = {
        "work_experience": [],
        "project_experience": [],
        "education_related": [],
        "general_skills": []
    }
    
    # Define context sections
    work_keywords = ['experience', 'employment', 'work', 'job', 'position', 'role', 'responsibilities']
    project_keywords = ['project', 'portfolio', 'built', 'developed', 'created', 'implemented']
    education_keywords = ['education', 'university', 'college', 'degree', 'course', 'studied', 'learned']
    
    text_lower = text.lower()
    
    for skill in skills:
        skill_lower = skill.lower()
        
        # Find context where skill appears
        skill_positions = [m.start() for m in re.finditer(re.escape(skill_lower), text_lower)]
        
        for pos in skill_positions:
            # Look at surrounding context (200 chars before and after)
            context_start = max(0, pos - 200)
            context_end = min(len(text), pos + 200)
            context = text_lower[context_start:context_end]
            
            # Categorize based on context
            if any(keyword in context for keyword in work_keywords):
                if skill not in skill_contexts["work_experience"]:
                    skill_contexts["work_experience"].append(skill)
            elif any(keyword in context for keyword in project_keywords):
                if skill not in skill_contexts["project_experience"]:
                    skill_contexts["project_experience"].append(skill)
            elif any(keyword in context for keyword in education_keywords):
                if skill not in skill_contexts["education_related"]:
                    skill_contexts["education_related"].append(skill)
            else:
                if skill not in skill_contexts["general_skills"]:
                    skill_contexts["general_skills"].append(skill)
    
    return skill_contexts

def _assess_professional_maturity(text: str, entities: Dict[str, Any]) -> Dict[str, Any]:
    """Assess professional maturity based on various indicators"""
    text_lower = text.lower()
    
    # Senior-level indicators
    senior_indicators = [
        'senior', 'lead', 'principal', 'architect', 'manager', 'director', 'head of',
        'chief', 'vp', 'vice president', 'cto', 'ceo', 'founder'
    ]
    
    # Mentorship indicators
    mentorship_indicators = [
        'mentor', 'mentored', 'mentoring', 'coached', 'training', 'onboarded',
        'guided', 'supervised', 'taught'
    ]
    
    # Strategic thinking indicators
    strategic_indicators = [
        'strategy', 'strategic', 'roadmap', 'vision', 'planning', 'architecture',
        'scalability', 'optimization', 'performance improvement'
    ]
    
    # Calculate scores
    senior_score = sum(1 for indicator in senior_indicators if indicator in text_lower)
    mentorship_score = sum(1 for indicator in mentorship_indicators if indicator in text_lower)
    strategic_score = sum(1 for indicator in strategic_indicators if indicator in text_lower)
    
    # Experience years (if available)
    experience_years = entities.get("experience_years", [])
    max_experience = max([int(year) for year in experience_years if year.isdigit()], default=0)
    
    # Professional maturity score (0-10)
    maturity_score = min(10, senior_score + mentorship_score + strategic_score + (max_experience // 2))
    
    return {
        "maturity_score": maturity_score,
        "senior_level_indicators": senior_score,
        "mentorship_indicators": mentorship_score,
        "strategic_thinking_indicators": strategic_score,
        "estimated_experience_years": max_experience
    }

def _extract_context_patterns(text: str) -> Dict[str, List[str]]:
    """Extract context patterns for better semantic understanding"""
    patterns = {
        "action_verbs": [],
        "quantified_achievements": [],
        "technology_stacks": [],
        "industry_keywords": []
    }
    
    # Action verbs commonly used in resumes
    action_verbs_pattern = r'\b(led|managed|developed|implemented|designed|built|created|optimized|improved|increased|reduced|delivered|achieved|coordinated|collaborated|established|initiated|streamlined|automated|modernized|transformed|scaled|migrated|integrated|deployed|maintained|monitored|analyzed|researched|presented|trained|supervised|mentored)\b'
    
    action_verbs = re.findall(action_verbs_pattern, text, re.IGNORECASE)
    patterns["action_verbs"] = list(set(action_verbs))
    
    # Quantified achievements
    quantified_pattern = r'(?:increased|improved|reduced|saved|generated|achieved|delivered|grew|optimized|enhanced).*?(?:\d+%|\$\d+|x\d+|\d+:\d+|\d+\s*(?:hours?|days?|weeks?|months?|years?))'
    
    quantified_achievements = re.findall(quantified_pattern, text, re.IGNORECASE)
    patterns["quantified_achievements"] = quantified_achievements
    
    # Technology stacks (grouped technologies)
    tech_stack_patterns = [
        r'(?:stack|technologies?):\s*([^.\n]+)',
        r'(?:built\s+with|using|technologies?\s+used):\s*([^.\n]+)',
        r'(?:frontend|backend|full[\s-]?stack):\s*([^.\n]+)'
    ]
    
    for pattern in tech_stack_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        patterns["technology_stacks"].extend(matches)
    
    # Industry-specific keywords
    industry_keywords = [
        'fintech', 'healthcare', 'e-commerce', 'saas', 'startup', 'enterprise',
        'agile', 'devops', 'microservices', 'api', 'machine learning', 'ai',
        'blockchain', 'cryptocurrency', 'iot', 'cloud', 'big data', 'analytics'
    ]
    
    found_industry_keywords = [keyword for keyword in industry_keywords if keyword in text.lower()]
    patterns["industry_keywords"] = found_industry_keywords
    
    return patterns

def extract_docs_to_chunks_and_records(paths: List[Path]) -> Tuple[List[str], List[Dict[str,Any]], List[Dict[str,Any]]]:
    """
    Enhanced extraction with comprehensive entity recognition and EDA
    
    Returns:
      chunks: list of strings with semantic meaning
      metas:  list of metadata dicts with comprehensive entities
      records: list of per-doc summaries with advanced analytics
    """
    paths = [p for p in paths if p.suffix.lower() in ALLOWED]
    if not paths:
        return [], [], []
    
    total_bytes = sum(p.stat().st_size for p in paths if p.exists())
    params = _adaptive_params(total_bytes, len(paths))
    CHUNK, OVER = params["chunk"], params["overlap"]

    chunks: List[str] = []
    metas: List[Dict[str,Any]] = []
    records: List[Dict[str,Any]] = []

    for p in paths:
        full_text = ""
        urls = []
        page2urls = {}
        
        # Parse document based on type
        if p.suffix.lower() == ".pdf":
            full_text, urls, page2urls = _parse_pdf(p)
        elif p.suffix.lower() in {".doc", ".docx"}:
            full_text, urls = _parse_docx(p)
            page2urls = {}

        if not full_text.strip():
            continue

        # Comprehensive entity extraction
        entities = _extract_comprehensive_entities(full_text)
        
        # Advanced EDA
        eda_results = _perform_advanced_eda(full_text, entities)
        
        # Extract social links
        socials = sorted(set([u for u in urls if SOCIAL.search(u)]))
        if not socials:
            socials = [u for u in URL.findall(full_text) if SOCIAL.search(u)]
        
        # Intelligent chunking with semantic boundaries
        doc_chunks = _split_overlap(full_text, CHUNK, OVER)

        # Create metadata for each chunk
        for i, ch in enumerate(doc_chunks):
            # Extract entities specific to this chunk
            chunk_entities = _extract_comprehensive_entities(ch)
            
            # Estimate page number
            page_guess = 1
            if page2urls:
                total_pages = max(page2urls.keys()) if page2urls else 1
                page_guess = min(total_pages, max(1, round((i+1) / max(1, len(doc_chunks)) * total_pages)))

            chunk_meta = {
                "file": p.name,
                "path": str(p),
                "page": page_guess,
                "chunk_index": i,
                "emails": chunk_entities["emails"][:5],
                "names": chunk_entities["names"][:5],
                "phones": chunk_entities["phones"][:3],
                "social_links": socials[:5],
                "skills": chunk_entities["skills"],
                "organizations": chunk_entities["organizations"][:5],
                "locations": chunk_entities["locations"][:5],
                "certifications": chunk_entities["certifications"][:5],
                "education": chunk_entities["education"][:3],
                "experience_years": chunk_entities["experience_years"],
                "char_count": len(ch),
                "word_count": len(ch.split()),
                "readability": chunk_entities.get("readability_stats", {})
            }
            metas.append(chunk_meta)
            chunks.append(ch)

        # Create comprehensive document record
        record = {
            "file": p.name,
            "path": str(p),
            "file_hash": file_sha256(p),
            "total_chunks": len(doc_chunks),
            "total_pages": max(page2urls.keys()) if page2urls else 1,
            "file_size": p.stat().st_size,
            "entities": entities,
            "eda_analysis": eda_results,
            "social_links": socials,
            "chunk_params": {"chunk_size": CHUNK, "overlap": OVER},
            "processing_metadata": {
                "spacy_available": _NLP is not None,
                "transformers_available": _NER_PIPELINE is not None,
                "textstat_available": _HAS_TEXTSTAT
            }
        }
        records.append(record)

    return chunks, metas, records
   
