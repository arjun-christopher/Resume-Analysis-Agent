# app/parsing_pymupdf.py - Advanced resume parsing using PyMuPDF with comprehensive information extraction
from __future__ import annotations

import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import hashlib
import re
import json
import time
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Dict, List, Tuple, Optional, Set
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

# Performance optimization libraries
try:
    import joblib
    from joblib import Parallel, delayed
    _HAS_JOBLIB = True
except ImportError:
    _HAS_JOBLIB = False
    Parallel = None
    delayed = None

try:
    import pickle
    _HAS_PICKLE = True
except ImportError:
    _HAS_PICKLE = False

try:
    from multiprocessing import cpu_count
    _CPU_COUNT = cpu_count()
except ImportError:
    _CPU_COUNT = 1

# Async I/O support
try:
    import asyncio
    import aiofiles
    _HAS_ASYNC = True
except ImportError:
    _HAS_ASYNC = False
    asyncio = None
    aiofiles = None

# Core libraries
import fitz  # PyMuPDF
import dateparser
import phonenumbers
from docx import Document as DocxDocument

# NLP and ML libraries
try:
    import spacy
    # Try to load transformer model first, fallback to standard
    try:
        _NLP = spacy.load("en_core_web_trf")
    except OSError:
        try:
            _NLP = spacy.load("en_core_web_sm")
        except OSError:
            _NLP = None
except ImportError:
    _NLP = None

try:
    from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
    # Load advanced NER model (with error handling for compatibility issues)
    try:
        _NER_PIPELINE = pipeline(
            "ner", 
            model="dbmdz/bert-large-cased-finetuned-conll03-english",
            aggregation_strategy="simple",
            device=-1  # Use CPU
        )
        _HAS_TRANSFORMERS = True
    except (RuntimeError, AttributeError, ImportError) as e:
        print(f"Transformers NER pipeline disabled due to compatibility issue: {e}")
        _NER_PIPELINE = None
        _HAS_TRANSFORMERS = False
except ImportError:
    _NER_PIPELINE = None
    _HAS_TRANSFORMERS = False

try:
    from sentence_transformers import SentenceTransformer
    try:
        _SENTENCE_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
        _HAS_SENTENCE_TRANSFORMERS = True
    except (RuntimeError, AttributeError, ImportError) as e:
        print(f"Sentence transformers disabled due to compatibility issue: {e}")
        _SENTENCE_MODEL = None
        _HAS_SENTENCE_TRANSFORMERS = False
except ImportError:
    _SENTENCE_MODEL = None
    _HAS_SENTENCE_TRANSFORMERS = False

try:
    from rapidfuzz import fuzz, process
    _HAS_RAPIDFUZZ = True
except ImportError:
    _HAS_RAPIDFUZZ = False

try:
    from flashtext import KeywordProcessor
    _HAS_FLASHTEXT = True
except ImportError:
    _HAS_FLASHTEXT = False

try:
    import textstat
    _HAS_TEXTSTAT = True
except ImportError:
    _HAS_TEXTSTAT = False

from extractors.certification_extractor import extract_certifications_comprehensive
from extractors.project_extractor import extract_projects_comprehensive
from extractors.achievements_extractor import extract_achievements_comprehensive
from extractors.activities_extractor import extract_activities_comprehensive
from extractors.publications_extractor import extract_publications_comprehensive
from extractors.experience_extractor import extract_experiences_comprehensive
from extractors.skills_extractor import extract_comprehensive_skills, SKILLS_DATABASE, _KEYWORD_PROCESSOR
from extractors.education_extractor import extract_education_comprehensive, extract_education_info
from extractors.hobbies_extractor import extract_hobbies_comprehensive
from extractors.languages_extractor import extract_languages_comprehensive
from extractors.summary_extractor import extract_summary_comprehensive

# ---------- Configuration ----------
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx"}
MAX_FILES_PER_SESSION = 50
MAX_TEXT_LENGTH = 1_000_000  # Limit text processing for performance

# Cache configuration
CACHE_DIR = Path("data/cache")
CACHE_DIR.mkdir(parents=True, exist_ok=True)
CACHE_FILE = CACHE_DIR / "processed_resumes.json"  # Legacy JSON cache
CACHE_JOBLIB_FILE = CACHE_DIR / "processed_resumes.pkl"  # Joblib cache with embeddings

# Pre-compiled regex patterns for efficiency
EMAIL_PATTERN = re.compile(
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    re.IGNORECASE
)

PHONE_PATTERN = re.compile(
    r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\b\d{10}\b',
    re.MULTILINE
)

URL_PATTERN = re.compile(
    r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
    re.IGNORECASE
)

# Enhanced hyperlink classification patterns (ordered from most specific to most general)
from collections import OrderedDict

HYPERLINK_PATTERNS = OrderedDict([
    ('email', re.compile(r'mailto:([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', re.IGNORECASE)),
    ('linkedin', re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/(?:in|pub|company)/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('github', re.compile(r'(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('twitter', re.compile(r'(?:https?://)?(?:www\.)?(?:twitter|x)\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('youtube', re.compile(r'(?:https?://)?(?:www\.)?(?:youtube\.com/(?:c/|channel/|user/)?|youtu\.be/)([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('kaggle', re.compile(r'(?:https?://)?(?:www\.)?kaggle\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('medium', re.compile(r'(?:https?://)?(?:www\.)?medium\.com/@?([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('stackoverflow', re.compile(r'(?:https?://)?(?:www\.)?stackoverflow\.com/users/(\d+/[A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('behance', re.compile(r'(?:https?://)?(?:www\.)?behance\.net/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('dribbble', re.compile(r'(?:https?://)?(?:www\.)?dribbble\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('instagram', re.compile(r'(?:https?://)?(?:www\.)?instagram\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    ('facebook', re.compile(r'(?:https?://)?(?:www\.)?facebook\.com/([A-Za-z0-9\-_.]+)', re.IGNORECASE)),
    # Portfolio pattern for specific domains
    ('portfolio', re.compile(r'(?:https?://)?(?:www\.)?([A-Za-z0-9\-_.]+\.(?:portfolio|website|dev|me|io))', re.IGNORECASE)),
    # Website pattern - most general, should be last
    ('website', re.compile(r'(?:https?://)?(?:www\.)?([A-Za-z0-9\-_.]+\.[A-Za-z]{2,})', re.IGNORECASE)),
])

# Text label patterns for hyperlinks with generic display text
TEXT_LABEL_PATTERNS = {
    'linkedin': [
        r'\blinkedin\b', r'\blinked\s*in\b', r'\bprofile\b', r'\bprofessional\s*profile\b',
        r'\bconnect\b', r'\bconnect\s*with\s*me\b', r'\bview\s*profile\b'
    ],
    'github': [
        r'\bgithub\b', r'\bgit\s*hub\b', r'\bcode\b', r'\brepository\b', r'\brepo\b',
        r'\bprojects\b', r'\bmy\s*code\b', r'\bview\s*code\b', r'\bgit\b'
    ],
    'twitter': [
        r'\btwitter\b', r'\bfollow\s*me\b', r'\btweet\b', r'\bfollow\b', r'\b@\w+\b',
        r'\bx\.com\b', r'\bx\b'
    ],
    'website': [
        r'\bwebsite\b', r'\bvisit\b', r'\bhomepage\b', r'\bweb\b', r'\bsite\b',
        r'\bclick\s*here\b', r'\bmore\s*info\b', r'\blearn\s*more\b', r'\bview\b'
    ],
    'portfolio': [
        r'\bportfolio\b', r'\bwork\b', r'\bmy\s*work\b', r'\bprojects\b', 
        r'\bshowcase\b', r'\bgallery\b', r'\bexamples\b'
    ],
    'medium': [
        r'\bmedium\b', r'\bblog\b', r'\barticles\b', r'\bwriting\b', r'\bposts\b',
        r'\bread\s*more\b', r'\bmy\s*blog\b'
    ],
    'email': [
        r'\bemail\b', r'\bcontact\b', r'\bmail\b', r'\bget\s*in\s*touch\b',
        r'\breach\s*out\b', r'\bmessage\b'
    ],
    'youtube': [
        r'\byoutube\b', r'\bvideo\b', r'\bvideos\b', r'\bchannel\b', r'\bwatch\b',
        r'\bsubscribe\b', r'\bmy\s*channel\b'
    ],
    'kaggle': [
        r'\bkaggle\b', r'\bdata\s*science\b', r'\bcompetitions\b', r'\bdatasets\b'
    ],
    'stackoverflow': [
        r'\bstack\s*overflow\b', r'\bso\b', r'\bstackoverflow\b', r'\breputation\b'
    ]
}

# Compile text label patterns for efficiency
COMPILED_TEXT_PATTERNS = {}
for platform, patterns in TEXT_LABEL_PATTERNS.items():
    COMPILED_TEXT_PATTERNS[platform] = [re.compile(pattern, re.IGNORECASE) for pattern in patterns]

# Legacy patterns for backward compatibility
SOCIAL_PATTERNS = {
    'linkedin': HYPERLINK_PATTERNS['linkedin'],
    'github': HYPERLINK_PATTERNS['github'],
    'twitter': HYPERLINK_PATTERNS['twitter'],
    'kaggle': HYPERLINK_PATTERNS['kaggle'],
    'medium': HYPERLINK_PATTERNS['medium'],
    'stackoverflow': HYPERLINK_PATTERNS['stackoverflow'],
}

DATE_PATTERNS = [
    re.compile(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b', re.IGNORECASE),
    re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'),
    re.compile(r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b'),
    re.compile(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', re.IGNORECASE),
]

# Name extraction patterns
NAME_PATTERNS = [
    re.compile(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]*\.?\s*)*[A-Z][a-z]+)$'),
    re.compile(r'^([A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+)$'),
    re.compile(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})$'),
]

# ---------- Skills and Education Extraction - Moved to separate modules ----------
# Skills extraction functions are now in extractors/skills_extractor.py
# Education extraction functions are now in extractors/education_extractor.py

# Skills extraction functions moved to extractors/skills_extractor.py

# Education extraction patterns and functions moved to extractors/education_extractor.py


# Education extraction functions moved to extractors/education_extractor.py

# ---------- Utility Functions ----------
def file_sha256(file_path: Path) -> str:
    """Calculate SHA256 hash of file"""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()

def load_cache() -> Dict[str, Any]:
    """Load cache of processed resumes (prefers joblib, falls back to JSON)"""
    # Try joblib cache first (includes embeddings)
    if _HAS_JOBLIB and CACHE_JOBLIB_FILE.exists():
        try:
            cache = joblib.load(CACHE_JOBLIB_FILE)
            print(f"Loaded cache from joblib: {len(cache)} resumes")
            return cache
        except Exception as e:
            print(f"Error loading joblib cache: {e}")
    
    # Fallback to JSON cache (legacy, no embeddings)
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, 'r') as f:
                cache = json.load(f)
                print(f"Loaded cache from JSON: {len(cache)} resumes (no embeddings)")
                return cache
        except Exception as e:
            print(f"Error loading JSON cache: {e}")
            return {}
    return {}

def save_cache(cache: Dict[str, Any]):
    """Save cache of processed resumes (uses joblib for efficiency)"""
    # Save with joblib for fast serialization with embeddings
    if _HAS_JOBLIB:
        try:
            joblib.dump(cache, CACHE_JOBLIB_FILE, compress=3)
            print(f"Saved cache with joblib: {len(cache)} resumes")
        except Exception as e:
            print(f"Error saving joblib cache: {e}")
    
    # Also save JSON for backward compatibility (without embeddings)
    try:
        # Remove embeddings from JSON cache to keep it lightweight
        json_cache = {}
        for resume_id, data in cache.items():
            json_data = data.copy()
            if 'chunks' in json_data:
                del json_data['chunks']
            if 'chunk_metadata' in json_data:
                # Remove embedding data from metadata
                json_data['chunk_metadata'] = [
                    {k: v for k, v in meta.items() if k != 'embedding'}
                    for meta in json_data['chunk_metadata']
                ]
            json_cache[resume_id] = json_data
        
        with open(CACHE_FILE, 'w') as f:
            json.dump(json_cache, f, indent=2)
    except Exception as e:
        print(f"Error saving JSON cache: {e}")

def is_resume_cached(resume_id: str, file_path: Path, cache: Dict[str, Any]) -> bool:
    """Check if resume is already cached and file hasn't changed"""
    if resume_id not in cache:
        return False
    
    cached_entry = cache[resume_id]
    cached_timestamp = cached_entry.get('file_modified')
    current_timestamp = file_path.stat().st_mtime
    
    # Check if file has been modified since caching
    return cached_timestamp == current_timestamp

def get_cached_resume(resume_id: str, cache: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """Retrieve cached resume data"""
    return cache.get(resume_id)

def cache_resume(resume_id: str, file_path: Path, chunks: List[str], 
                 chunk_metadata: List[Dict[str, Any]], document_record: Dict[str, Any], 
                 cache: Dict[str, Any]):
    """Cache processed resume data with complete chunks and embeddings"""
    cache[resume_id] = {
        'resume_id': resume_id,
        'resume_name': file_path.name,
        'file_path': str(file_path),
        'file_modified': file_path.stat().st_mtime,
        'chunk_count': len(chunks),
        'chunks': chunks,  # Store actual chunks for instant retrieval
        'chunk_metadata': chunk_metadata,  # Store metadata with embeddings
        'document_record': document_record,
        'processing_timestamp': time.time()
    }

def normalize_phone_number(phone: str) -> Optional[str]:
    """Normalize phone number using phonenumbers library"""
    try:
        parsed = phonenumbers.parse(phone, "US")  # Default to US
        if phonenumbers.is_valid_number(parsed):
            return phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
    except:
        pass
    return phone

def parse_date_with_dateparser(date_str: str) -> Optional[str]:
    """Parse date string using dateparser"""
    try:
        parsed_date = dateparser.parse(date_str)
        if parsed_date:
            return parsed_date.strftime("%Y-%m-%d")
    except:
        pass
    return date_str

def _is_plausible_match(url: str, platform: str) -> bool:
    """Check if a URL could plausibly belong to a platform based on domain or common patterns"""
    url_lower = url.lower()
    
    # Common domain patterns that would exclude certain platforms
    if platform == 'linkedin' and 'linkedin' not in url_lower:
        # Could still be valid if it's a redirect or short URL
        return not any(domain in url_lower for domain in ['github.com', 'twitter.com', 'instagram.com', 'facebook.com'])
    elif platform == 'github' and 'github' not in url_lower:
        return not any(domain in url_lower for domain in ['linkedin.com', 'twitter.com', 'instagram.com', 'facebook.com'])
    elif platform == 'twitter' and not any(x in url_lower for x in ['twitter', 'x.com']):
        return not any(domain in url_lower for domain in ['linkedin.com', 'github.com', 'instagram.com', 'facebook.com'])
    
    # If URL contains the platform name, it's very likely a match
    if platform.lower() in url_lower:
        return True
    
    # For generic labels like "website", "visit", accept most URLs
    if platform in ['website', 'portfolio']:
        return True
    
    # If no obvious conflicts, allow the match
    return True

def _extract_username_from_url(url: str, platform: str) -> Optional[str]:
    """Try to extract username from URL for a given platform"""
    if platform in HYPERLINK_PATTERNS:
        match = HYPERLINK_PATTERNS[platform].search(url)
        if match and match.groups():
            return match.group(1)
    
    # Fallback: try to extract username from common URL patterns
    try:
        # Remove protocol and www
        clean_url = re.sub(r'^https?://(www\.)?', '', url)
        # Split by / and take relevant parts
        parts = clean_url.split('/')
        if len(parts) >= 2:
            return parts[1]  # Usually the username part
    except:
        pass
    
    return None

def extract_and_classify_hyperlinks(doc: fitz.Document) -> Dict[str, Any]:
    """
    Extract hyperlinks from PDF using PyMuPDF and classify them by type
    
    Args:
        doc: PyMuPDF document object
        
    Returns:
        Dictionary with classified hyperlinks by type and page
    """
    
    hyperlinks_data = {
        'by_type': {
            'email': [],
            'linkedin': [],
            'github': [],
            'twitter': [],
            'kaggle': [],
            'medium': [],
            'stackoverflow': [],
            'behance': [],
            'dribbble': [],
            'portfolio': [],
            'youtube': [],
            'instagram': [],
            'facebook': [],
            'website': [],
            'other': []
        },
        'by_page': {},
        'all_links': [],
        'statistics': {
            'total_links': 0,
            'pages_with_links': 0,
            'most_common_type': None
        }
    }
    
    pages_with_links = 0
    
    try:
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_links = []
            
            # Extract hyperlinks using the method you provided
            links = page.get_links()
            
            for link in links:
                if "uri" in link:
                    uri = link["uri"]
                    if uri and isinstance(uri, str):
                        page_links.append(uri)
                        
                        # Extract display text from the link rectangle area
                        display_text = ""
                        if "from" in link:
                            try:
                                # Get the rectangle coordinates
                                rect = link["from"]
                                # Extract text from the rectangle area
                                text_instances = page.get_text("dict", clip=rect)
                                if text_instances and "blocks" in text_instances:
                                    for block in text_instances["blocks"]:
                                        if "lines" in block:
                                            for line in block["lines"]:
                                                if "spans" in line:
                                                    for span in line["spans"]:
                                                        if "text" in span:
                                                            display_text += span["text"] + " "
                                display_text = display_text.strip()
                            except:
                                pass
                        
                        hyperlinks_data['all_links'].append({
                            'url': uri,
                            'display_text': display_text,
                            'page': page_num + 1,
                            'type': None  # Will be classified below
                        })
            
            # Also extract hyperlinks from annotations (alternative method)
            try:
                for annot in page.annots():
                    if annot.type[1] == 'Link':  # Link annotation
                        link_data = annot.info.get('content', '') or annot.info.get('title', '')
                        if link_data and 'http' in link_data:
                            page_links.append(link_data)
                            
                            # Try to get display text from annotation
                            display_text = annot.info.get('title', '') or ""
                            
                            hyperlinks_data['all_links'].append({
                                'url': link_data,
                                'display_text': display_text,
                                'page': page_num + 1,
                                'type': None
                            })
            except:
                pass  # Annotations might not be accessible
            
            if page_links:
                pages_with_links += 1
                hyperlinks_data['by_page'][page_num + 1] = page_links
        
        # Classify all extracted hyperlinks using both URL and text patterns
        for link_entry in hyperlinks_data['all_links']:
            url = link_entry['url']
            display_text = link_entry.get('display_text', '').lower().strip()
            classified = False
            username = None
            
            # First, try to classify by URL pattern (most reliable)
            for link_type, pattern in HYPERLINK_PATTERNS.items():
                match = pattern.search(url)
                if match:
                    username = match.group(1) if match.groups() else None
                    hyperlinks_data['by_type'][link_type].append({
                        'url': url,
                        'display_text': link_entry.get('display_text', ''),
                        'page': link_entry['page'],
                        'username': username,
                        'classification_method': 'url_pattern'
                    })
                    link_entry['type'] = link_type
                    classified = True
                    break
            
            # If not classified by URL, try text label patterns
            if not classified and display_text:
                for platform, text_patterns in COMPILED_TEXT_PATTERNS.items():
                    for text_pattern in text_patterns:
                        if text_pattern.search(display_text):
                            # Double-check that the URL could plausibly match this platform
                            if _is_plausible_match(url, platform):
                                hyperlinks_data['by_type'][platform].append({
                                    'url': url,
                                    'display_text': link_entry.get('display_text', ''),
                                    'page': link_entry['page'],
                                    'username': _extract_username_from_url(url, platform),
                                    'classification_method': 'text_pattern'
                                })
                                link_entry['type'] = platform
                                classified = True
                                break
                    if classified:
                        break
            
            # If still not classified, mark as other
            if not classified:
                hyperlinks_data['by_type']['other'].append({
                    'url': url,
                    'display_text': link_entry.get('display_text', ''),
                    'page': link_entry['page'],
                    'username': None,
                    'classification_method': 'unclassified'
                })
                link_entry['type'] = 'other'
        
        # Calculate statistics
        hyperlinks_data['statistics']['total_links'] = len(hyperlinks_data['all_links'])
        hyperlinks_data['statistics']['pages_with_links'] = pages_with_links
        
        # Find most common link type
        type_counts = {
            link_type: len(links) 
            for link_type, links in hyperlinks_data['by_type'].items() 
            if links
        }
        
        if type_counts:
            most_common_type = max(type_counts.items(), key=lambda x: x[1])
            hyperlinks_data['statistics']['most_common_type'] = {
                'type': most_common_type[0],
                'count': most_common_type[1]
            }
    
    except Exception as e:
        print(f"Error extracting hyperlinks: {e}")
    
    return hyperlinks_data

# ---------- PyMuPDF-based PDF Extraction ----------
def extract_pdf_with_pymupdf(file_path: Path) -> Tuple[str, List[str], Dict[str, Any]]:
    """
    Extract text, links, and metadata from PDF using PyMuPDF (fitz) with enhanced hyperlink classification
    
    Returns:
        tuple: (full_text, urls, metadata)
    """
    full_text = []
    urls = []
    page_metadata = {}
    
    try:
        doc = fitz.open(str(file_path))
        
        # Enhanced hyperlink extraction and classification
        hyperlinks_data = extract_and_classify_hyperlinks(doc)
        
        # Document metadata
        metadata = {
            "page_count": doc.page_count,
            "metadata": doc.metadata,
            "is_encrypted": doc.is_encrypted,
            "is_pdf": doc.is_pdf,
            "hyperlinks": hyperlinks_data,
            "page_details": {}
        }
        
        # Collect all URLs for backward compatibility
        urls = [link['url'] for link in hyperlinks_data['all_links']]
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Extract text with layout information
            text = page.get_text()
            if text.strip():
                full_text.append(text)
            
            # Get page-specific hyperlinks from our classification
            page_links = hyperlinks_data['by_page'].get(page_num + 1, [])
            
            # Extract images metadata (for future OCR if needed)
            images = page.get_images()
            
            # Extract text blocks with position information
            text_blocks = page.get_text("dict")
            
            page_metadata[page_num] = {
                "links": page_links,
                "link_count": len(page_links),
                "image_count": len(images),
                "text_blocks": len(text_blocks.get("blocks", [])),
                "page_rect": list(page.rect)
            }
        
        doc.close()
        
        # Extract comprehensive skills from the full text
        full_text_str = "\n".join(full_text)
        comprehensive_skills = extract_comprehensive_skills(full_text_str)
        
        metadata["page_details"] = page_metadata
        metadata["skills"] = comprehensive_skills
        
        return full_text_str, urls, metadata
        
    except Exception as e:
        print(f"Error extracting PDF {file_path}: {e}")
        return "", [], {}

def extract_docx_content(file_path: Path) -> Tuple[str, List[str]]:
    """Extract text and hyperlinks from DOCX file"""
    try:
        doc = DocxDocument(str(file_path))
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract table content
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    tables.append(" | ".join(cells))
        
        # Extract hyperlinks
        urls = []
        rels = getattr(doc.part, "rels", {})
        for rel in rels.values():
            try:
                if "hyperlink" in rel.reltype and rel.target_ref:
                    urls.append(str(rel.target_ref))
            except:
                continue
        
        full_text = "\n".join(paragraphs + tables)
        return full_text, urls
        
    except Exception as e:
        print(f"Error extracting DOCX {file_path}: {e}")
        return "", []

# ---------- Information Extraction Functions ----------
def extract_emails(text: str) -> List[str]:
    """Extract and validate email addresses"""
    emails = EMAIL_PATTERN.findall(text)
    # Additional validation
    valid_emails = []
    for email in emails:
        if "@" in email and "." in email.split("@")[-1]:
            valid_emails.append(email.lower())
    return list(set(valid_emails))

def extract_phone_numbers(text: str) -> List[str]:
    """Extract and normalize phone numbers"""
    phones = PHONE_PATTERN.findall(text)
    normalized_phones = []
    
    for phone in phones:
        # Clean the phone number
        cleaned = re.sub(r'[^\d+]', '', phone)
        if len(cleaned) >= 10:  # Valid phone number length
            normalized = normalize_phone_number(phone)
            if normalized:
                normalized_phones.append(normalized)
    
    return list(set(normalized_phones))

def extract_urls_and_social_links(text: str, hyperlinks_data: Optional[Dict] = None) -> Dict[str, List[str]]:
    """
    Extract URLs and categorize social media links from text and hyperlinks
    
    Args:
        text: Text content to search for URLs
        hyperlinks_data: Optional classified hyperlinks from PDF extraction
    """
    # Extract URLs from text using regex
    text_urls = URL_PATTERN.findall(text)
    
    social_links = {platform: [] for platform in SOCIAL_PATTERNS.keys()}
    social_links['other_urls'] = []
    social_links['email'] = []
    social_links['portfolio'] = []
    social_links['website'] = []
    
    # Process URLs found in text
    for url in text_urls:
        categorized = False
        for platform, pattern in SOCIAL_PATTERNS.items():
            match = pattern.search(url)
            if match:
                social_links[platform].append(url)
                categorized = True
                break
        
        if not categorized:
            # Check for other types
            if EMAIL_PATTERN.search(url):
                social_links['email'].append(url)
            elif any(domain in url.lower() for domain in ['.portfolio', '.me', '.dev', '.website']):
                social_links['portfolio'].append(url)
            else:
                social_links['other_urls'].append(url)
    
    # If we have classified hyperlinks data from PDF, merge it
    if hyperlinks_data and isinstance(hyperlinks_data, dict):
        hyperlink_types = hyperlinks_data.get('by_type', {})
        
        for link_type, links in hyperlink_types.items():
            if links and link_type in social_links:
                # Extract just the URLs from the link objects
                urls_to_add = [link['url'] for link in links if isinstance(link, dict) and 'url' in link]
                social_links[link_type].extend(urls_to_add)
            elif links and link_type == 'other':
                urls_to_add = [link['url'] for link in links if isinstance(link, dict) and 'url' in link]
                social_links['other_urls'].extend(urls_to_add)
    
    # Remove duplicates and clean up
    for key in social_links:
        social_links[key] = list(set(social_links[key]))
    
    return social_links

def extract_dates(text: str) -> List[Dict[str, str]]:
    """Extract and normalize dates"""
    dates = []
    
    for pattern in DATE_PATTERNS:
        matches = pattern.findall(text)
        for match in matches:
            normalized_date = parse_date_with_dateparser(match)
            dates.append({
                "original": match,
                "normalized": normalized_date
            })
    
    return dates

def format_candidate_name(name: str) -> str:
    """
    Format candidate name with proper capitalization, handling initials and special cases
    
    Examples:
        "ARJUN CHRISTOPHER" -> "Arjun Christopher"
        "arjun christopher" -> "Arjun Christopher"
        "JOHN SMITH JR" -> "John Smith Jr"
        "A.R.RAHMAN" -> "A.R. Rahman"
        "MC DONALD" -> "Mc Donald"
        "O'BRIEN" -> "O'Brien"
        "Jean-Paul SARTRE" -> "Jean Paul Sartre" (hyphen removed)
        "arjun_christopher" -> "Arjun Christopher" (underscore removed)
    """
    if not name or not isinstance(name, str):
        return name
    
    # First, replace hyphens and underscores with spaces
    name = name.replace('-', ' ').replace('_', ' ')
    
    # Clean up extra whitespace
    name = ' '.join(name.split())
    
    # Handle edge case: initials stuck to name like "A.R.RAHMAN" -> "A.R. RAHMAN"
    # But preserve "A.R." together as one unit
    # Look for pattern where multiple dots exist followed by uppercase letters
    name = re.sub(r'([A-Z]\.[A-Z]\.?)([A-Z][A-Za-z]+)', r'\1 \2', name)
    
    def format_word(word: str) -> str:
        """Format a single word with proper capitalization"""
        if not word:
            return word
        
        # Handle initials (single letter or letter with period)
        if len(word) <= 2 and word[0].isalpha():
            # Single letter initial or "A." format
            if len(word) == 1:
                return word.upper() + '.'
            elif len(word) == 2 and word[1] == '.':
                return word[0].upper() + '.'
        
        # Handle initials with dots (e.g., "A.R." or "A.R" or "a.r.rahman" or "DR.")
        if '.' in word:
            # Check if it's a title like "Dr." or suffix
            word_lower = word.lower().rstrip('.')
            title_suffix_patterns = {
                'jr': 'Jr', 'sr': 'Sr', 
                'ii': 'II', 'iii': 'III', 'iv': 'IV', 'v': 'V',
                'phd': 'PhD', 'md': 'MD', 'esq': 'Esq',
                'dds': 'DDS', 'jd': 'JD', 'dr': 'Dr', 'mr': 'Mr',
                'mrs': 'Mrs', 'ms': 'Ms', 'prof': 'Prof'
            }
            if word_lower in title_suffix_patterns:
                return title_suffix_patterns[word_lower] + '.'
            
            # Handle multi-part initials like "A.R." or "a.r.something"
            parts = word.split('.')
            formatted_parts = []
            for i, p in enumerate(parts):
                if not p:  # Empty part (from consecutive dots or trailing dot)
                    continue
                if len(p) == 1 and p.isalpha():
                    # Single letter - make it uppercase (initial)
                    formatted_parts.append(p.upper())
                elif len(p) == 2 and p.isalpha():
                    # Two letters together like "AR" in "A.R.NAME"
                    formatted_parts.append(p.upper())
                else:
                    # Regular word part
                    formatted_parts.append(p.capitalize())
            
            result = '.'.join(formatted_parts)
            # Add trailing period if original had it and result doesn't
            if word.endswith('.') and not result.endswith('.'):
                result += '.'
            return result
        
        # Handle apostrophes (e.g., "O'Brien", "D'Angelo")
        if "'" in word:
            parts = word.split("'")
            if len(parts) == 2 and len(parts[0]) <= 2:
                # Handles O'Brien, D'Angelo
                return parts[0].capitalize() + "'" + parts[1].capitalize()
            return "'".join(part.capitalize() for part in parts)
        
        # Handle Mc/Mac prefixes (e.g., "McDonald", "MacLeod")
        if word.lower().startswith('mc') and len(word) > 2:
            return 'Mc' + word[2:].capitalize()
        
        # For "MacDonald" as single word, just capitalize normally
        if word.lower().startswith('mac') and len(word) > 3:
            return word.capitalize()
        
        # Handle common suffixes (Jr, Sr, II, III, IV, PhD, MD, etc.) without dots
        suffix_patterns = {
            'jr': 'Jr', 'sr': 'Sr', 
            'ii': 'II', 'iii': 'III', 'iv': 'IV', 'v': 'V',
            'phd': 'PhD', 'md': 'MD', 'esq': 'Esq',
            'dds': 'DDS', 'jd': 'JD', 'dr': 'Dr'
        }
        if word.lower() in suffix_patterns:
            return suffix_patterns[word.lower()]
        
        # Standard capitalization for regular words
        return word.capitalize()
    
    # Split name into words and format each
    words = name.split()
    formatted_words = [format_word(word) for word in words]
    
    return ' '.join(formatted_words)

def extract_names_advanced(text: str) -> List[str]:
    """Advanced name extraction using multiple strategies with confidence scoring"""
    candidate_names = {}  # Store names with confidence scores
    lines = text.split('\n')
    
    # Enhanced skip keywords to avoid false positives
    skip_keywords = [
        'resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary', 
        'contact', 'email', 'phone', 'address', 'skills', 'experience',
        'education', 'objective', 'portfolio', 'projects', 'work',
        'professional', 'personal', 'references', 'certification',
        'achievement', 'publication', 'awards', 'languages', 'hobbies'
    ]
    
    # Words to exclude from being considered as names
    exclude_words = [
        'the', 'and', 'for', 'with', 'from', 'about', 'university',
        'college', 'school', 'company', 'corporation', 'inc', 'llc',
        'engineer', 'developer', 'manager', 'analyst', 'designer',
        'consultant', 'specialist', 'director', 'senior', 'junior'
    ]
    
    def calculate_name_confidence(name: str, position: int, context: str) -> float:
        """Calculate confidence score for a potential name"""
        score = 0.0
        words = name.split()
        
        # Position bonus (earlier in document = higher confidence)
        if position < 3:
            score += 50
        elif position < 10:
            score += 30
        elif position < 20:
            score += 10
        
        # Length bonus (2-4 words typical for names)
        if 2 <= len(words) <= 4:
            score += 30
        elif len(words) == 1:
            score -= 20
        
        # All words capitalized
        if all(word[0].isupper() for word in words if word):
            score += 20
        
        # No numbers or special characters except hyphens and apostrophes
        if not any(char.isdigit() or char in '!@#$%^&*()_+=[]{}|;:,.<>?/' for char in name):
            score += 15
        
        # Context clues
        context_lower = context.lower()
        if any(keyword in context_lower for keyword in skip_keywords):
            score -= 30
        
        # Check if words are common name parts
        common_prefixes = ['mr', 'mrs', 'ms', 'dr', 'prof']
        common_suffixes = ['jr', 'sr', 'ii', 'iii', 'iv', 'phd', 'md']
        
        for word in words:
            word_lower = word.lower().rstrip('.,')
            if word_lower in common_prefixes:
                score += 10
            elif word_lower in common_suffixes:
                score += 5
            elif word_lower in exclude_words:
                score -= 40
        
        return score
    
    # Strategy 1: Look in first lines with enhanced pattern matching
    for i, line in enumerate(lines[:25]):
        original_line = line
        line = line.strip()
        
        if len(line) < 3 or len(line) > 150:
            continue
        
        # Remove common prefixes
        line_clean = re.sub(r'^(Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)\s*', '', line, flags=re.IGNORECASE)
        
        # Try to match name patterns
        # Pattern 1: Simple capitalized words at start of line
        if i < 5 and line_clean and line_clean[0].isupper():
            # Check if line looks like a name (not a sentence)
            if not line.endswith(('.', '!', '?')) and len(line.split()) <= 5:
                words = line_clean.split()
                # Filter out words that are likely not names
                name_words = [w for w in words if w and w[0].isupper() and not any(char.isdigit() for char in w)]
                
                if 2 <= len(name_words) <= 4:
                    potential_name = ' '.join(name_words)
                    # Clean up any trailing punctuation
                    potential_name = re.sub(r'[,;:]$', '', potential_name)
                    
                    confidence = calculate_name_confidence(potential_name, i, original_line)
                    if confidence > 20:  # Threshold
                        if potential_name not in candidate_names or candidate_names[potential_name] < confidence:
                            candidate_names[potential_name] = confidence
        
        # Pattern 2: Name patterns with regex
        for pattern in NAME_PATTERNS:
            match = pattern.match(line_clean)
            if match:
                potential_name = match.group(1).strip()
                # Remove any trailing punctuation
                potential_name = re.sub(r'[,;:]$', '', potential_name)
                words = potential_name.split()
                
                if 2 <= len(words) <= 4 and all(len(w) >= 2 for w in words):
                    confidence = calculate_name_confidence(potential_name, i, original_line)
                    if confidence > 20:
                        if potential_name not in candidate_names or candidate_names[potential_name] < confidence:
                            candidate_names[potential_name] = confidence
    
    # Strategy 2: Extract from email usernames (lower confidence)
    emails = extract_emails(text)
    for email in emails:
        username = email.split('@')[0]
        # Handle various email username formats
        if '.' in username:
            parts = username.split('.')[:2]  # Take first two parts
            if all(part.isalpha() and len(part) >= 2 for part in parts):
                name = ' '.join(part.title() for part in parts)
                confidence = 15  # Lower confidence from email
                if name not in candidate_names or candidate_names[name] < confidence:
                    candidate_names[name] = confidence
        elif '_' in username:
            parts = username.split('_')[:2]
            if all(part.isalpha() and len(part) >= 2 for part in parts):
                name = ' '.join(part.title() for part in parts)
                confidence = 15
                if name not in candidate_names or candidate_names[name] < confidence:
                    candidate_names[name] = confidence
    
    # Strategy 3: Use spaCy NER if available (high confidence for PERSON entities in first 5000 chars)
    if _NLP:
        try:
            doc = _NLP(text[:5000])  # Limit for performance
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.strip()) > 3:
                    name = ent.text.strip()
                    # Calculate position-based confidence
                    position_in_text = text[:5000].find(name)
                    line_number = text[:position_in_text].count('\n') if position_in_text >= 0 else 999
                    
                    confidence = 25  # Base confidence for NER
                    if line_number < 10:
                        confidence += 25
                    elif line_number < 20:
                        confidence += 15
                    
                    if name not in candidate_names or candidate_names[name] < confidence:
                        candidate_names[name] = confidence
        except Exception as e:
            print(f"SpaCy NER error: {e}")
    
    # Strategy 4: Use transformer-based NER if available
    if _NER_PIPELINE:
        try:
            # Process first 2000 characters for efficiency
            text_sample = text[:2000]
            ner_results = _NER_PIPELINE(text_sample)
            
            for entity in ner_results:
                if entity.get('entity_group') == 'PER' or 'PER' in entity.get('entity', ''):
                    name = entity.get('word', '').strip()
                    # Clean up tokenization artifacts
                    name = name.replace('##', '').replace('Ġ', ' ').strip()
                    
                    if len(name) > 3:
                        # Position-based confidence
                        start_pos = entity.get('start', 1000)
                        line_number = text_sample[:start_pos].count('\n')
                        
                        confidence = 30  # Higher confidence for transformer
                        if line_number < 5:
                            confidence += 30
                        elif line_number < 15:
                            confidence += 15
                        
                        if name not in candidate_names or candidate_names[name] < confidence:
                            candidate_names[name] = confidence
        except Exception as e:
            print(f"Transformer NER error: {e}")
    
    # Sort by confidence and return top names
    sorted_names = sorted(candidate_names.items(), key=lambda x: x[1], reverse=True)
    
    # Return names with confidence > threshold, formatted properly
    result = [format_candidate_name(name) for name, conf in sorted_names if conf > 20]
    
    # If we found names, return top 3 max; if none, return empty list
    return result[:3] if result else []

def extract_skills_flashtext(text: str) -> Dict[str, List[str]]:
    """Extract skills using FlashText for fast keyword matching"""
    if _KEYWORD_PROCESSOR:
        found_skills = _KEYWORD_PROCESSOR.extract_keywords(text.lower())
        
        # Categorize skills
        categorized_skills = {category: [] for category in SKILLS_DATABASE.keys()}
        
        for skill in found_skills:
            for category, skills_set in SKILLS_DATABASE.items():
                if skill in skills_set:
                    categorized_skills[category].append(skill)
        
        return categorized_skills
    
    return {}

def extract_skills_fuzzy(text: str, threshold: int = 80) -> Dict[str, List[str]]:
    """Extract skills using fuzzy matching with RapidFuzz"""
    if not _HAS_RAPIDFUZZ:
        return {}
    
    words = re.findall(r'\b\w+\b', text.lower())
    text_phrases = []
    
    # Create 1-3 word phrases
    for i in range(len(words)):
        text_phrases.append(words[i])
        if i < len(words) - 1:
            text_phrases.append(f"{words[i]} {words[i+1]}")
        if i < len(words) - 2:
            text_phrases.append(f"{words[i]} {words[i+1]} {words[i+2]}")
    
    found_skills = {category: [] for category in SKILLS_DATABASE.keys()}
    
    for category, skills_set in SKILLS_DATABASE.items():
        for skill in skills_set:
            # Find best match
            best_match = process.extractOne(
                skill, 
                text_phrases, 
                scorer=fuzz.partial_ratio,
                score_cutoff=threshold
            )
            
            if best_match:
                found_skills[category].append(skill)
    
    return found_skills

def extract_organizations_and_locations(text: str) -> Tuple[List[str], List[str]]:
    """Extract organizations and locations using NER"""
    organizations = []
    locations = []
    
    if _NLP:
        doc = _NLP(text[:10000])  # Limit for performance
        
        for ent in doc.ents:
            if ent.label_ == "ORG":
                organizations.append(ent.text.strip())
            elif ent.label_ in ["GPE", "LOC"]:
                locations.append(ent.text.strip())
    
    # Also use transformer-based NER if available
    if _NER_PIPELINE:
        try:
            # Process in chunks
            chunk_size = 500
            for i in range(0, min(len(text), 5000), chunk_size):
                chunk = text[i:i+chunk_size]
                entities = _NER_PIPELINE(chunk)
                
                for entity in entities:
                    if entity['entity_group'] == 'ORG':
                        organizations.append(entity['word'].strip())
                    elif entity['entity_group'] == 'LOC':
                        locations.append(entity['word'].strip())
        except Exception as e:
            print(f"Error in transformer NER: {e}")
    
    return list(set(organizations)), list(set(locations))

def calculate_readability_stats(text: str) -> Dict[str, float]:
    """Calculate readability statistics"""
    if not _HAS_TEXTSTAT:
        return {}
    
    try:
        stats = {
            "flesch_reading_ease": textstat.flesch_reading_ease(text),
            "flesch_kincaid_grade": textstat.flesch_kincaid_grade(text),
            "automated_readability_index": textstat.automated_readability_index(text),
            "coleman_liau_index": textstat.coleman_liau_index(text),
            "word_count": len(text.split()),
            "sentence_count": textstat.sentence_count(text),
            "avg_sentence_length": textstat.avg_sentence_length(text),
            "syllable_count": textstat.syllable_count(text),
        }
        return stats
    except Exception as e:
        print(f"Error calculating readability: {e}")
        return {}

def extract_experience_years(text: str) -> List[int]:
    """Extract years of experience mentions"""
    experience_patterns = [
        r'(\d+)[\s]*(?:\+|\-)?[\s]*(?:years?|yrs?|year)[\s]*(?:of[\s]*)?(?:experience|exp)',
        r'(?:experience|exp)[\s]*(?:of[\s]*)?(\d+)[\s]*(?:\+|\-)?[\s]*(?:years?|yrs?|year)',
        r'(\d+)[\s]*(?:\+|\-)?[\s]*(?:years?|yrs?)[\s]*(?:in|with|of)',
    ]
    
    years = []
    for pattern in experience_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if match.isdigit():
                year_val = int(match)
                if 0 <= year_val <= 50:  # Reasonable range
                    years.append(year_val)
    
    return sorted(set(years))

def extract_education_info(text: str) -> List[Dict[str, str]]:
    """Extract education information"""
    education_patterns = [
        r'(?:Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)\s*(?:of|in|degree)?\s*([A-Za-z\s]{3,50})',
        r'([A-Za-z\s]{3,50})\s*(?:degree|diploma|certificate)',
        r'(?:University|College|Institute)\s+of\s+([A-Za-z\s]{3,50})',
    ]
    
    education = []
    for pattern in education_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            education.append({
                "degree": match.strip(),
                "type": "inferred"
            })
    
    return education

# ---------- Main Extraction Function ----------
def extract_comprehensive_entities(text: str, hyperlinks_data: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Comprehensive entity extraction using multiple techniques
    
    Args:
        text: Text content to analyze
        hyperlinks_data: Optional classified hyperlinks from PDF extraction
    """
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]
    
    # Basic extractions
    emails = extract_emails(text)
    phones = extract_phone_numbers(text)
    social_data = extract_urls_and_social_links(text, hyperlinks_data)
    dates = extract_dates(text)
    names = extract_names_advanced(text)
    
    # Enhanced comprehensive skills extraction
    comprehensive_skills = extract_comprehensive_skills(text)
    
    # Legacy skills extraction for backward compatibility
    legacy_skills = extract_skills_flashtext(text)
    if not any(legacy_skills.values()) and _HAS_RAPIDFUZZ:
        legacy_skills = extract_skills_fuzzy(text)
    
    # NER extractions
    organizations, locations = extract_organizations_and_locations(text)
    
    # Additional extractions
    experience_years = extract_experience_years(text)
    education = extract_education_comprehensive(text)
    readability = calculate_readability_stats(text)
    certifications = extract_certifications_comprehensive(text)
    projects = extract_projects_comprehensive(text)
    achievements = extract_achievements_comprehensive(text)
    activities = extract_activities_comprehensive(text)
    publications = extract_publications_comprehensive(text)
    experiences = extract_experiences_comprehensive(text)
    hobbies = extract_hobbies_comprehensive(text)
    languages = extract_languages_comprehensive(text)
    professional_summary = extract_summary_comprehensive(text)

    # Compile results
    entities = {
        "names": names,
        "emails": emails,
        "phones": phones,
        "social_links": social_data,
        "dates": dates,
        "skills": comprehensive_skills,
        "legacy_skills": legacy_skills,
        "organizations": organizations,
        "locations": locations,
        "experience_years": experience_years,
        "education": education,
        "readability_stats": readability,
        "certifications": certifications,
        "projects": projects,
        "achievements": achievements,
        "activities": activities,
        "publications": publications,
        "experiences": experiences,
        # New extractors for commonly missed sections
        "hobbies": hobbies,
        "languages": languages,
        "professional_summary": professional_summary,
        "processing_info": {
            "text_length": len(text),
            "spacy_available": _NLP is not None,
            "transformers_available": _HAS_TRANSFORMERS,
            "flashtext_available": _HAS_FLASHTEXT,
            "rapidfuzz_available": _HAS_RAPIDFUZZ,
            "textstat_available": _HAS_TEXTSTAT,
        }
    }
    
    return entities

# ---------- Semantic Similarity (if sentence-transformers available) ----------
def calculate_semantic_similarity(text1: str, text2: str) -> float:
    """Calculate semantic similarity between two texts"""
    if _HAS_SENTENCE_TRANSFORMERS and _SENTENCE_MODEL:
        try:
            embeddings = _SENTENCE_MODEL.encode([text1, text2])
            similarity = float(_SENTENCE_MODEL.similarity(embeddings[0], embeddings[1]))
            return similarity
        except Exception as e:
            print(f"Error calculating similarity: {e}")
    
    return 0.0

def extract_semantic_chunks(text: str, chunk_size: int = 512, overlap: int = 50) -> List[Dict[str, Any]]:
    """
    Extract semantically meaningful chunks with embeddings using batch processing
    
    Args:
        text: Text to chunk
        chunk_size: Target chunk size in words
        overlap: Overlap size in words
    
    Returns:
        List of chunk dictionaries with text, metadata, and embeddings
    """
    chunks = []
    
    # Split text into sentences first
    sentences = re.split(r'[.!?]+', text)
    
    current_chunk = []
    current_length = 0
    chunk_texts = []  # Collect all chunk texts for batch embedding
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        sentence_length = len(sentence.split())
        
        if current_length + sentence_length > chunk_size and current_chunk:
            # Create chunk
            chunk_text = '. '.join(current_chunk)
            chunk_info = {
                "text": chunk_text,
                "word_count": len(chunk_text.split()),
                "sentence_count": len(current_chunk)
            }
            
            chunks.append(chunk_info)
            chunk_texts.append(chunk_text)
            
            # Handle overlap
            overlap_sentences = current_chunk[-overlap//50:] if overlap > 0 else []
            current_chunk = overlap_sentences + [sentence]
            current_length = sum(len(s.split()) for s in current_chunk)
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    
    # Add final chunk
    if current_chunk:
        chunk_text = '. '.join(current_chunk)
        chunk_info = {
            "text": chunk_text,
            "word_count": len(chunk_text.split()),
            "sentence_count": len(current_chunk)
        }
        
        chunks.append(chunk_info)
        chunk_texts.append(chunk_text)
    
    # Batch embed all chunks at once for better performance
    if _HAS_SENTENCE_TRANSFORMERS and _SENTENCE_MODEL and chunk_texts:
        try:
            # Encode all chunks in one batch call (much faster than one-by-one)
            batch_embeddings = _SENTENCE_MODEL.encode(
                chunk_texts,
                batch_size=64,  # Process 64 chunks at a time
                show_progress_bar=False,
                convert_to_numpy=True
            )
            
            # Assign embeddings to chunks
            for i, embedding in enumerate(batch_embeddings):
                if i < len(chunks):
                    chunks[i]["embedding"] = embedding.tolist()
        except Exception as e:
            print(f"Warning: Batch embedding failed: {e}")
    
    return chunks

# ---------- Single File Processing (for parallel execution) ----------
async def async_read_file_bytes(file_path: Path) -> bytes:
    """
    Asynchronously read file bytes for non-blocking I/O
    
    Args:
        file_path: Path to the file
    
    Returns:
        File content as bytes
    """
    if _HAS_ASYNC:
        async with aiofiles.open(file_path, 'rb') as f:
            return await f.read()
    else:
        # Fallback to synchronous read
        with open(file_path, 'rb') as f:
            return f.read()

async def async_process_file_batch(file_paths: List[Path], cache: Dict[str, Any]) -> List[Optional[Tuple]]:
    """
    Asynchronously process a batch of files with non-blocking I/O
    
    Args:
        file_paths: List of file paths to process
        cache: Cache dictionary
    
    Returns:
        List of results from processing each file
    """
    # Read all files asynchronously
    tasks = [async_read_file_bytes(fp) for fp in file_paths]
    file_contents = await asyncio.gather(*tasks)
    
    # Process files (CPU-bound operations still synchronous)
    results = []
    for file_path, content in zip(file_paths, file_contents):
        # Process the file using the content we just read
        result = process_single_file(file_path, cache)
        results.append(result)
    
    return results

def process_single_file(file_path: Path, cache: Dict[str, Any]) -> Optional[Tuple[str, List[str], List[Dict[str, Any]], Dict[str, Any]]]:
    """
    Process a single resume file (used for parallel processing)
    
    Args:
        file_path: Path to the file
        cache: Cache dictionary
    
    Returns:
        Tuple of (resume_id, chunks, chunk_metadata, document_record) or None if cached/failed
    """
    if not file_path.exists():
        return None
    
    # Generate unique resume ID using file hash
    resume_id = file_sha256(file_path)
    resume_name = file_path.name
    
    # Check cache first
    if is_resume_cached(resume_id, file_path, cache):
        print(f"✓ [Cached] {file_path.name}")
        cached_data = get_cached_resume(resume_id, cache)
        if cached_data and 'chunks' in cached_data and 'chunk_metadata' in cached_data:
            return (
                resume_id,
                cached_data['chunks'],
                cached_data['chunk_metadata'],
                cached_data['document_record']
            )
    
    print(f"⚙ [Processing] {file_path.name}")
    
    # Extract content based on file type
    try:
        if file_path.suffix.lower() == ".pdf":
            full_text, urls, pdf_metadata = extract_pdf_with_pymupdf(file_path)
            hyperlinks_data = pdf_metadata.get('hyperlinks', {})
        elif file_path.suffix.lower() in {".doc", ".docx"}:
            full_text, urls = extract_docx_content(file_path)
            pdf_metadata = {}
            hyperlinks_data = {}
        else:
            return None
        
        if not full_text.strip():
            print(f"⚠ No text extracted from {file_path.name}")
            return None
        
        # Extract comprehensive entities with hyperlink data
        entities = extract_comprehensive_entities(full_text, hyperlinks_data)
        
        # Extract candidate name with high priority
        candidate_names = extract_names_advanced(full_text)
        candidate_name = candidate_names[0] if candidate_names else None
        
        # If no name found through advanced extraction, try to get from entities
        if not candidate_name and entities.get('names'):
            candidate_name = format_candidate_name(entities['names'][0])
        
        # Fallback: Use filename (remove extension and clean) with proper formatting
        if not candidate_name:
            filename_name = file_path.stem.replace('_', ' ').replace('-', ' ')
            candidate_name = format_candidate_name(filename_name)
        
        # Create semantic chunks
        semantic_chunks = extract_semantic_chunks(full_text)
        
        chunks = []
        chunk_metadata = []
        
        # Process each chunk
        for i, chunk_info in enumerate(semantic_chunks):
            chunk_text = chunk_info["text"]
            
            # Extract chunk-specific entities (pass hyperlinks for consistency)
            chunk_entities = extract_comprehensive_entities(chunk_text, hyperlinks_data)
            
            # Create chunk metadata with resume identification
            chunk_meta = {
                "resume_id": resume_id,
                "resume_name": resume_name,
                "candidate_name": candidate_name,
                "file": file_path.name,
                "path": str(file_path),
                "chunk_index": i,
                "total_chunks": len(semantic_chunks),
                "char_count": len(chunk_text),
                "word_count": chunk_info["word_count"],
                "sentence_count": chunk_info["sentence_count"],
                "entities": chunk_entities,
                "urls": urls,
                "has_embedding": "embedding" in chunk_info
            }
            
            if "embedding" in chunk_info:
                chunk_meta["embedding"] = chunk_info["embedding"]
            
            chunks.append(chunk_text)
            chunk_metadata.append(chunk_meta)
        
        # Create document record with resume identification
        document_record = {
            "resume_id": resume_id,
            "resume_name": resume_name,
            "candidate_name": candidate_name,
            "file": file_path.name,
            "path": str(file_path),
            "file_hash": resume_id,
            "file_size": file_path.stat().st_size,
            "total_chunks": len(semantic_chunks),
            "entities": entities,
            "urls": urls,
            "pdf_metadata": pdf_metadata,
            "processing_timestamp": time.time(),
        }
        
        return (resume_id, chunks, chunk_metadata, document_record)
    
    except Exception as e:
        print(f"✗ Error processing {file_path.name}: {e}")
        return None

# ---------- Main Processing Function ----------
def process_documents_with_pymupdf(
    file_paths: List[Path], 
    use_parallel: bool = True, 
    n_jobs: int = -1,
    use_async_io: bool = True
) -> Tuple[List[str], List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    Process documents using PyMuPDF with comprehensive information extraction, caching, parallel processing, and async I/O
    
    Args:
        file_paths: List of file paths to process
        use_parallel: Whether to use parallel processing (default: True)
        n_jobs: Number of parallel jobs (-1 for all CPUs, default: -1)
        use_async_io: Whether to use async I/O for file reading (default: True)
        
    Returns:
        tuple: (chunks, chunk_metadata, document_records)
    """
    
    file_paths = [p for p in file_paths if p.suffix.lower() in ALLOWED_EXTENSIONS]
    if not file_paths:
        return [], [], []
    
    chunks = []
    chunk_metadata = []
    document_records = []
    
    # Load cache
    cache = load_cache()
    cache_updated = False
    
    # Determine if parallel processing should be used
    use_parallel = use_parallel and _HAS_JOBLIB and len(file_paths) > 1
    
    # Async I/O mode (for large number of files)
    if use_async_io and _HAS_ASYNC and len(file_paths) > 5:
        print(f"Processing {len(file_paths)} files with async I/O...")
        
        # Run async batch processing
        if asyncio.get_event_loop().is_running():
            # If already in an event loop, create a new one
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            results = loop.run_until_complete(async_process_file_batch(file_paths, cache))
            loop.close()
        else:
            results = asyncio.run(async_process_file_batch(file_paths, cache))
        
        # Collect results
        for result in results:
            if result is not None:
                resume_id, file_chunks, file_metadata, document_record = result
                chunks.extend(file_chunks)
                chunk_metadata.extend(file_metadata)
                document_records.append(document_record)
                
                # Update cache
                cache_resume(resume_id, Path(document_record['path']), 
                           file_chunks, file_metadata, document_record, cache)
                cache_updated = True
    
    elif use_parallel:
        # Parallel processing mode
        # Calculate number of jobs
        if n_jobs == -1:
            n_jobs = min(_CPU_COUNT, len(file_paths))
        else:
            n_jobs = min(n_jobs, len(file_paths))
        
        print(f"Processing {len(file_paths)} files in parallel using {n_jobs} workers...")
        
        # Process files in parallel
        results = Parallel(n_jobs=n_jobs, backend='loky', verbose=0)(
            delayed(process_single_file)(file_path, cache) for file_path in file_paths
        )
        
        # Collect results
        for result in results:
            if result is not None:
                resume_id, file_chunks, file_metadata, document_record = result
                chunks.extend(file_chunks)
                chunk_metadata.extend(file_metadata)
                document_records.append(document_record)
                
                # Update cache
                cache_resume(resume_id, Path(document_record['path']), 
                           file_chunks, file_metadata, document_record, cache)
                cache_updated = True
    else:
        # Sequential processing (fallback or single file)
        print(f"Processing {len(file_paths)} files sequentially...")
        
        for file_path in file_paths:
            result = process_single_file(file_path, cache)
            if result is not None:
                resume_id, file_chunks, file_metadata, document_record = result
                chunks.extend(file_chunks)
                chunk_metadata.extend(file_metadata)
                document_records.append(document_record)
                
                # Update cache
                cache_resume(resume_id, file_path, file_chunks, file_metadata, document_record, cache)
                cache_updated = True
    
    # Save cache if updated
    if cache_updated:
        save_cache(cache)
        print(f"✓ Cache updated with {len(document_records)} resume(s)")
    
    print(f"✓ Processed {len(document_records)} documents into {len(chunks)} chunks")
    return chunks, chunk_metadata, document_records

# ---------- Export Function for Compatibility ----------
def extract_docs_to_chunks_and_records(paths: List[Path]) -> Tuple[List[str], List[Dict[str,Any]], List[Dict[str,Any]]]:
    """
    Compatibility function with the original parsing interface
    """
    return process_documents_with_pymupdf(paths)
